from pathlib import Path
import re
import sys

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC = next((ROOT / "docs").glob("NYT2539-2016*.docx"))
OUT_MD = ROOT / "docs" / "nyt2539_database_mapping.md"
OUT_PRESETS = ROOT / "backend" / "app" / "db" / "dictionary_presets.py"


def clean(text: str | None) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u3000", " ")).strip()


def rows_of(table) -> list[list[str]]:
    return [[clean(cell.text) for cell in row.cells] for row in table.rows]


def md_escape(value: object) -> str:
    return str(value or "").replace("|", "\\|").replace("\n", "<br>")


def table_line(values: list[object]) -> str:
    return "| " + " | ".join(md_escape(value) for value in values) + " |"


def py_str(value: str | int | None) -> str:
    return repr(value)


doc = Document(str(DOC))

b_tables = [
    ("B.1", "CBDKXX", "承包地块信息"),
    ("B.2", "FBF", "发包方"),
    ("B.3", "CBF", "承包方"),
    ("B.4", "CBF_JTCY", "承包方家庭成员"),
    ("B.5", "CBHT", "承包合同"),
    ("B.6", "LZHT", "流转合同"),
    ("B.7", "QSLYZLFJ", "权属来源资料附件"),
    ("B.8", "CBJYQZDJB", "承包经营权证登记簿"),
    ("B.9", "CBJYQZ", "承包经营权证"),
    ("B.10", "CBJYQZ_QZBF", "权证补发"),
    ("B.11", "CBJYQZ_QZHF", "权证换发"),
    ("B.12", "CBJYQZ_QZZX", "权证注销"),
]

std_tables = {}
for offset, (no, code, name) in enumerate(b_tables):
    rows = rows_of(doc.tables[18 + offset])
    fields = []
    for row in rows[1:]:
        if len(row) < 3 or not row[2]:
            continue
        if not re.fullmatch(r"[A-Z0-9_]+", row[2]):
            continue
        fields.append(
            {
                "name": row[1] if len(row) > 1 else "",
                "code": row[2],
                "type": row[3] if len(row) > 3 else "",
                "length": row[4] if len(row) > 4 else "",
                "scale": row[5] if len(row) > 5 else "",
                "domain": row[6] if len(row) > 6 else "",
                "constraint": row[7] if len(row) > 7 else "",
                "note": row[8] if len(row) > 8 else "",
            }
        )
    std_tables[code] = {"no": no, "code": code, "name": name, "fields": fields}

sys.path.insert(0, str(ROOT / "backend"))
import app.models  # noqa: E402,F401
from app.db.base import Base  # noqa: E402

current_tables = {
    table.name: [column.name for column in table.columns]
    for table in sorted(Base.metadata.tables.values(), key=lambda item: item.name)
}

manual_mapping = {
    "CBDKXX": {
        "tables": ["survey_cbdkxx_base", "survey_cbdkxx_result"],
        "status": "已落库（调查基表/成果表）",
        "note": "按批次保留原始快照和调查成果，字段基本按规范代码小写存储，并增加批次、来源追踪、成果状态等业务字段。",
    },
    "FBF": {
        "tables": ["fbf", "survey_fbf_base", "survey_fbf_result", "issuers"],
        "status": "已落库（历史表 + 调查表 + 主数据视图）",
        "note": "fbf 保留规范字段；survey_fbf_* 用于调查批次；issuers 是系统业务侧发包方管理表，字段语义映射但命名更业务化。",
    },
    "CBF": {
        "tables": ["survey_cbf_base", "survey_cbf_result"],
        "status": "已落库（调查基表/成果表）",
        "note": "旧 cbf 表会迁移到 survey_cbf_* 后删除；当前以调查批次模型承载承包方。",
    },
    "CBF_JTCY": {
        "tables": ["survey_cbf_jtcy_base", "survey_cbf_jtcy_result"],
        "status": "已落库（调查基表/成果表）",
        "note": "成果表扩展了成员状态、政策依据、权益处置等业务字段。",
    },
    "CBHT": {
        "tables": ["cbht"],
        "status": "已落库（规范表）",
        "note": "字段按规范代码小写保存，并增加租户和区域字段。",
    },
    "LZHT": {
        "tables": [],
        "status": "未单独建表",
        "note": "当前未发现流转合同专表；如需管理流转合同，可按 B.6 新增。",
    },
    "QSLYZLFJ": {
        "tables": ["survey_attachments", "request_case_attachments"],
        "status": "部分承载",
        "note": "系统用通用附件表承载调查/流程附件，未按 QSLYZLFJ 字段一比一建表。",
    },
    "CBJYQZDJB": {
        "tables": ["request_cases"],
        "status": "流程侧部分承载",
        "note": "登记簿未单独建表；证书登记业务目前主要进入流程申请和档案。",
    },
    "CBJYQZ": {
        "tables": ["request_cases"],
        "status": "流程侧部分承载",
        "note": "权证发放信息未单独建表。",
    },
    "CBJYQZ_QZBF": {
        "tables": ["request_cases"],
        "status": "流程侧部分承载",
        "note": "补发作为业务类型/流程处理，未按规范表单独持久化全部字段。",
    },
    "CBJYQZ_QZHF": {
        "tables": ["request_cases"],
        "status": "流程侧部分承载",
        "note": "换发作为业务类型/流程处理，未按规范表单独持久化全部字段。",
    },
    "CBJYQZ_QZZX": {
        "tables": ["request_cases"],
        "status": "流程侧部分承载",
        "note": "注销作为业务类型/流程处理，未按规范表单独持久化全部字段。",
    },
}

system_tables = {
    "tenants": "租户/县域数据隔离",
    "regions": "行政区划与区域树",
    "users": "用户账号",
    "roles": "角色",
    "permissions": "权限点",
    "user_region_permissions": "用户可操作区域",
    "dictionary_items": "系统字典项",
    "map_layers": "地图图层配置",
    "data_import_batches": "数据导入批次",
    "data_import_files": "导入文件",
    "data_import_rows": "导入行明细",
    "survey_batches": "承包方/地块调查批次",
    "survey_contractor_tasks": "调查任务",
    "survey_change_records": "调查变更记录",
    "survey_change_diffs": "字段级变更差异",
    "request_cases": "业务申请/流程实例",
    "request_case_participants": "流程办理记录",
    "request_case_attachments": "流程附件",
    "request_attachment_templates": "流程附件目录模板",
    "request_workflow_mappings": "业务类型到流程定义映射",
    "workflow_definition_versions": "流程定义版本",
}

title_map = {}
for paragraph in doc.paragraphs:
    text = clean(paragraph.text)
    match = re.search(r"C\.(\d+)\s*(.+)", text)
    if match and text.startswith("表"):
        title_map[int(match.group(1))] = f"C.{int(match.group(1))} {match.group(2)}"

dict_type_slug = {
    1: "nyt2539_c01_control_point_type_grade",
    2: "nyt2539_c02_marker_stone_type",
    3: "nyt2539_c03_marker_type",
    4: "nyt2539_c04_boundary_type",
    5: "nyt2539_c05_boundary_property",
    6: "nyt2539_c06_ownership_property",
    7: "nyt2539_c07_parcel_category",
    8: "nyt2539_c08_land_grade",
    9: "nyt2539_c09_land_use",
    10: "nyt2539_c10_right_acquire_method",
    11: "nyt2539_c11_boundary_point_type",
    12: "nyt2539_c12_boundary_marker_type",
    13: "nyt2539_c13_boundary_line_category",
    14: "nyt2539_c14_boundary_line_position",
    15: "nyt2539_c15_id_document_type",
    16: "nyt2539_c16_contractor_type",
    17: "nyt2539_c17_gender",
    18: "nyt2539_c18_member_remark",
    19: "nyt2539_c19_yes_no",
}

appendix_c_dicts = []
appendix_c_by_type = []
for n in range(1, 20):
    rows = rows_of(doc.tables[29 + n])
    dict_type = dict_type_slug[n]
    dict_name = title_map.get(n, f"C.{n}")
    entries = []
    for order, row in enumerate(rows[1:], start=1):
        if not row or not row[0]:
            continue
        code = row[0]
        if code.startswith("注") or code.startswith("说明"):
            continue
        cols = [value for value in row[1:] if value]
        name = " / ".join(cols) if cols else code
        remark = None
        if n == 1 and len(row) >= 3 and row[1]:
            remark = f"控制点类型：{row[1]}"
        item = (dict_type, dict_name, code, name, order * 10, remark)
        appendix_c_dicts.append(item)
        entries.append(item)
    appendix_c_by_type.append((dict_type, dict_name, entries))

preset_lines = ["NYT2539_APPENDIX_C_DICTIONARY_ITEMS = ["]
for item in appendix_c_dicts:
    preset_lines.append("    (" + ", ".join(py_str(value) for value in item) + "),")
preset_lines.append("]\n")
OUT_PRESETS.write_text("\n".join(preset_lines), encoding="utf-8")

md = []
md.append("# NY/T 2539-2016 与系统数据库结构对照")
md.append("")
md.append(f"- 来源文档：`docs/{DOC.name}`")
md.append("- 目的：说明规范中的权属数据表在当前系统数据库中的落点，便于后续开发、导入、查询和扩展。")
md.append("- 约定：规范字段代码在系统表中通常使用小写列名；调查类数据采用 `*_base` 保存原始/基准快照，`*_result` 保存调查成果。")
md.append("")
md.append("## 1. 总体表对照")
md.append("")
md.append(table_line(["规范表", "规范含义", "系统表", "状态", "说明"]))
md.append("|---|---|---|---|---|")
for code in [item[1] for item in b_tables]:
    info = std_tables[code]
    mapping = manual_mapping[code]
    md.append(
        table_line(
            [
                f"{info['no']} `{code}`",
                info["name"],
                ", ".join(f"`{table}`" for table in mapping["tables"]) or "未建表",
                mapping["status"],
                mapping["note"],
            ]
        )
    )
md.append("")
md.append("## 2. 当前数据库表分层")
md.append("")
md.append(table_line(["类别", "系统表", "用途"]))
md.append("|---|---|---|")
for table, purpose in system_tables.items():
    if table not in current_tables:
        continue
    category = "业务支撑"
    if table.startswith("survey_") or table in {"fbf", "cbht", "issuers"}:
        category = "权属/调查数据"
    elif table.startswith("request_") or table == "request_cases":
        category = "流程业务"
    elif table.startswith("data_import"):
        category = "导入追踪"
    elif table in {"tenants", "regions", "users", "roles", "permissions", "user_region_permissions", "dictionary_items"}:
        category = "系统管理"
    md.append(table_line([category, f"`{table}`", purpose]))
md.append("")
md.append("## 3. 字段级对照")
md.append("")
for code in [item[1] for item in b_tables]:
    info = std_tables[code]
    mapping = manual_mapping[code]
    target_cols = {}
    for table in mapping["tables"]:
        for column in current_tables.get(table, []):
            target_cols.setdefault(column.lower(), []).append(table)
    matched = sum(1 for field in info["fields"] if field["code"].lower() in target_cols)
    md.append(f"### {info['no']} `{code}` {info['name']}")
    md.append("")
    md.append(f"- 系统落点：{', '.join(f'`{table}`' for table in mapping['tables']) or '未单独建表'}")
    md.append(f"- 字段覆盖：{matched}/{len(info['fields'])} 个规范字段可按同名小写列直接对应。")
    md.append(f"- 说明：{mapping['note']}")
    md.append("")
    md.append(table_line(["规范字段", "代码", "类型", "值域/约束", "系统列落点"]))
    md.append("|---|---|---|---|---|")
    for field in info["fields"]:
        key = field["code"].lower()
        targets = target_cols.get(key, [])
        landing = ", ".join(f"`{table}.{key}`" for table in targets) if targets else "未发现同名列"
        md.append(
            table_line(
                [
                    field["name"],
                    f"`{field['code']}`",
                    " ".join(part for part in [field["type"], field["length"], field["scale"]] if part),
                    " / ".join(part for part in [field["domain"], field["constraint"]] if part),
                    landing,
                ]
            )
        )
    md.append("")
md.append("## 4. NY/T 2539 附录 C 字典落库")
md.append("")
md.append("附录 C 已整理为 `dictionary_items` 的初始化数据，字典类型采用 `nyt2539_cXX_...` 命名。业务页面可通过 `/api/v1/dictionaries/options/{dictType}` 读取，并由前端 `useDictionary` 自动缓存。")
md.append("")
md.append(table_line(["附录表", "dict_type", "条目数", "示例"]))
md.append("|---|---|---:|---|")
for dict_type, dict_name, entries in appendix_c_by_type:
    sample = ", ".join(f"{item[2]}={item[3]}" for item in entries[:3])
    md.append(table_line([dict_name, f"`{dict_type}`", len(entries), sample]))
md.append("")
md.append("## 5. 后续扩展建议")
md.append("")
md.append("- 若需要完整承载流转合同、权证登记簿、权证补换发/注销，应按 B.6-B.12 增建专表或在流程归档时固化结构化字段。")
md.append("- 对已落库的规范字段，导入程序应优先使用规范代码小写列，避免再新增中文或业务别名列。")
md.append("- 系统扩展字段应保留在规范字段之后，并在本对照文档中记录来源和用途。")
md.append("")
OUT_MD.write_text("\n".join(md), encoding="utf-8")

print(f"wrote {OUT_MD}")
print(f"wrote {OUT_PRESETS}")
print(f"appendix C entries: {len(appendix_c_dicts)}")
